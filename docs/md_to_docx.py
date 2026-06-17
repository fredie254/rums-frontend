"""
Convert all RUMS Markdown documentation files to Word (.docx) format.
Usage: python md_to_docx.py
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_ROOT = os.path.dirname(os.path.abspath(__file__))

MD_FILES = [
    ("business/BRD.md",                    "business/BRD.docx"),
    ("business/FRS.md",                    "business/FRS.docx"),
    ("business/URS.md",                    "business/URS.docx"),
    ("technical/system-architecture.md",   "technical/system-architecture.docx"),
    ("technical/database-design.md",       "technical/database-design.docx"),
    ("technical/api-documentation.md",     "technical/api-documentation.docx"),
    ("technical/security-architecture.md", "technical/security-architecture.docx"),
    ("technical/integration-specifications.md", "technical/integration-specifications.docx"),
    ("ux/wireframes.md",                   "ux/wireframes.docx"),
    ("ux/mockups.md",                      "ux/mockups.docx"),
    ("ux/user-journey-maps.md",            "ux/user-journey-maps.docx"),
]

# ── Colour constants ────────────────────────────────────────────────────────
COLOUR_HEADING1  = RGBColor(0x0d, 0x6e, 0xfd)   # Bootstrap primary blue
COLOUR_HEADING2  = RGBColor(0x19, 0x87, 0x54)   # Bootstrap success green
COLOUR_HEADING3  = RGBColor(0x49, 0x52, 0x57)   # Dark grey
COLOUR_CODE_BG   = RGBColor(0xf8, 0xf9, 0xfa)
COLOUR_CODE_FG   = RGBColor(0xd6, 0x33, 0x84)   # Pink/magenta for inline code
COLOUR_HR        = RGBColor(0xde, 0xe2, 0xe6)

# ── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour="F2F2F2"):
    """Set a table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run()
    run.add_break()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),  'single')
    bottom.set(qn('w:sz'),   '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DEE2E6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def apply_inline_markup(run_parent, text):
    """
    Parse **bold**, *italic*, `code`, and plain text within a paragraph.
    run_parent is a docx Paragraph object.
    """
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = run_parent.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = run_parent.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            r = run_parent.add_run(part[1:-1])
            r.font.name  = 'Courier New'
            r.font.size  = Pt(9)
            r.font.color.rgb = COLOUR_CODE_FG
        else:
            if part:
                run_parent.add_run(part)


def style_heading(para, level):
    """Apply custom colour + spacing to a heading paragraph."""
    colour = {1: COLOUR_HEADING1, 2: COLOUR_HEADING2, 3: COLOUR_HEADING3}.get(level, COLOUR_HEADING3)
    for run in para.runs:
        run.font.color.rgb = colour
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 8)
    para.paragraph_format.space_after  = Pt(4)


def add_code_block(doc, code_text):
    """Add a shaded code block paragraph with monospace font."""
    # Split into lines, add each as its own paragraph inside a table for shading
    lines = code_text.splitlines()
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F2F4F6')
    # Clear default empty paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for i, line in enumerate(lines):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x2d, 0x31, 0x39)
    doc.add_paragraph()   # spacer after block


def parse_table(doc, lines):
    """Parse a Markdown table (| col | col |) and add to doc."""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue   # separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    # Normalise row widths
    rows = [r + [''] * (col_count - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            apply_inline_markup(p, cell_text)
            if r_idx == 0:
                set_cell_bg(cell, '0D6EFD')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                    run.bold = True

    doc.add_paragraph()   # spacer


# ── Main converter ───────────────────────────────────────────────────────────

def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default body style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    lines = content.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Fenced code block ────────────────────────────────────────────────
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, '\n'.join(code_lines))
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────────────
        if re.match(r'^\s*---+\s*$', line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = min(len(m.group(1)), 4)
            text  = m.group(2).strip()
            # Strip inline anchor links like [text](#anchor)
            text  = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            para  = doc.add_heading(text, level=level)
            style_heading(para, level)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.strip().startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            parse_table(doc, table_lines)
            continue

        # ── Unordered list ────────────────────────────────────────────────────
        if re.match(r'^\s*[-*]\s+', line):
            para = doc.add_paragraph(style='List Bullet')
            text = re.sub(r'^\s*[-*]\s+', '', line)
            apply_inline_markup(para, text)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # ── Ordered list ──────────────────────────────────────────────────────
        if re.match(r'^\s*\d+\.\s+', line):
            para = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            apply_inline_markup(para, text)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.strip().startswith('>'):
            text = re.sub(r'^\s*>\s?', '', line)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent   = Cm(1.0)
            para.paragraph_format.space_before  = Pt(2)
            para.paragraph_format.space_after   = Pt(2)
            r = para.add_run(text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
            i += 1
            continue

        # ── Bold metadata lines (e.g. **Version:** 1.0) ───────────────────────
        # ── Regular paragraph / blank line ────────────────────────────────────
        stripped = line.strip()
        if stripped == '':
            # Only add space if previous para isn't a code block or heading
            i += 1
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        apply_inline_markup(para, stripped)
        i += 1

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"  OK  {os.path.relpath(docx_path, DOCS_ROOT)}")


# ── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("Converting RUMS documentation to Word (.docx)...\n")
    errors = []
    for md_rel, docx_rel in MD_FILES:
        md_path   = os.path.join(DOCS_ROOT, md_rel)
        docx_path = os.path.join(DOCS_ROOT, docx_rel)
        try:
            convert(md_path, docx_path)
        except Exception as e:
            print(f"  FAIL  {md_rel}: {e}")
            errors.append((md_rel, str(e)))

    print(f"\nDone. {len(MD_FILES) - len(errors)}/{len(MD_FILES)} files converted successfully.")
    if errors:
        print("Errors:")
        for f, e in errors:
            print(f"  {f}: {e}")
